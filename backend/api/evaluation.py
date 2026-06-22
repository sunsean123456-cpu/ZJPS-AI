"""
评审 API 路由
"""
import os
from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime

from database import get_db
from models.db_models import User, Project, Document, AIScore, ProjectStatus, AuditLog
from schemas import (
    EvaluationStartRequest, EvaluationResultResponse,
    AIScoreResponse, PreCheckResponse, MessageResponse
)
from services.auth_service import get_current_user, require_role
from services.evaluation_engine import evaluation_engine
from services.archive_agent import archive_agent
from services.document_parser import parser

router = APIRouter(prefix="/api/evaluation", tags=["评审"])


@router.get("/stream/{project_id}")
async def stream_evaluation(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    SSE 流式评审
    实时推送评审步骤：文档解析 → 要素提取 → 标准比对 → 评分生成 → 报告汇总
    """
    from fastapi.responses import StreamingResponse
    import asyncio
    import json
    
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    async def event_stream():
        """生成 SSE 事件流 - 真实评审流程"""
        from services.evaluation_engine import evaluation_engine
        
        steps = [
            {"id": 1, "name": "文档解析", "icon": "📄", "phase": "parse"},
            {"id": 2, "name": "要素提取", "icon": "🔍", "phase": "extract"},
            {"id": 3, "name": "标准比对", "icon": "⚖️", "phase": "compare"},
            {"id": 4, "name": "评分生成", "icon": "📊", "phase": "score"},
            {"id": 5, "name": "报告汇总", "icon": "✅", "phase": "report"},
        ]
        
        total_steps = len(steps)
        
        # 发送开始事件
        yield f"data: {json.dumps({'type': 'start', 'total_steps': total_steps, 'project_name': project.name})}\n\n"
        
        try:
            # 执行真实评审
            result = await evaluation_engine.run_stream_evaluation(project_id, db)
            
            # 逐步骤推送进度
            for i, step in enumerate(steps, 1):
                yield f"data: {json.dumps({'type': 'step_start', 'step': i, 'name': step['name'], 'icon': step['icon']})}\n\n"
                
                # 推送该阶段的进度
                phase = step['phase']
                if phase in result.get('progress', {}):
                    phase_progress = result['progress'][phase]
                    for progress in phase_progress:
                        yield f"data: {json.dumps({'type': 'step_progress', 'step': i, 'progress': progress})}\n\n"
                        await asyncio.sleep(0.3)
                
                msg = step['name'] + '完成'
                yield f"data: {json.dumps({'type': 'step_complete', 'step': i, 'message': msg})}\n\n"
            
            # 发送完成事件
            yield f"data: {json.dumps({'type': 'complete', 'total_score': result.get('total_score', 0), 'grade': result.get('grade', 'C')})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
    
    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",  # Nginx 不缓冲
        }
    )


@router.post("/start", response_model=dict)
async def start_evaluation(
    request: EvaluationStartRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    启动 AI 评审
    流程：本地预检 → 文档解析 → LLM评分 → 结果存储
    """
    project = db.query(Project).filter(Project.id == request.project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    # 权限检查
    if current_user.role.value != "admin" and project.enterprise_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权操作此项目")
    
    # 本地预检（零流量）
    pre_check = archive_agent.local_pre_check(project, db)
    if not pre_check["passed"]:
        return {
            "success": False,
            "stage": "pre_check",
            "message": pre_check["summary"],
            "pre_check": pre_check
        }
    
    # 更新状态为评审中
    project.status = ProjectStatus.reviewing
    project.current_step = 4
    db.commit()
    
    # 记录日志
    log = AuditLog(
        user_id=current_user.id,
        action="start_evaluation",
        module="evaluation",
        detail=f"启动项目评审: {project.name}"
    )
    db.add(log)
    db.commit()
    
    # 后台执行评审
    background_tasks.add_task(
        _run_evaluation_task,
        project_id=project.id,
        user_id=current_user.id
    )
    
    return {
        "success": True,
        "stage": "reviewing",
        "message": f"材料检查通过（完整性 {pre_check['completeness']:.0f}%），AI 评审已启动",
        "project_id": project.id
    }


@router.get("/result/{project_id}")
async def get_evaluation_result(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取评审结果"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    # 权限检查
    if current_user.role.value != "admin" and project.enterprise_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权查看此项目")
    
    if project.status not in [ProjectStatus.completed, ProjectStatus.sandbox_triggered]:
        return {
            "success": False,
            "message": f"评审尚未完成，当前状态: {project.status.value}"
        }
    
    # 获取评分详情
    scores = db.query(AIScore).filter(AIScore.project_id == project_id).all()
    dimensions = _build_score_dimensions(scores)
    
    # 生成建议
    suggestions = []
    for dim in dimensions:
        ratio = dim["score"] / dim["max_score"] if dim["max_score"] > 0 else 0
        if ratio < 0.7:
            suggestions.append(f"{dim['name']}得分较低（{dim['score']:.1f}/{dim['max_score']:.0f}），建议加强")
    
    return {
        "success": True,
        "project_id": project_id,
        "total_score": project.total_score,
        "grade": project.grade,
        "dimensions": dimensions,
        "summary": f"项目总得分 {project.total_score:.1f} 分，等级: {project.grade}",
        "suggestions": suggestions,
        "sandbox_recommended": project.sandbox_triggered,
        "completed_at": project.completed_at.isoformat() if project.completed_at else None
    }


@router.get("/scores/{project_id}", response_model=List[AIScoreResponse])
async def get_evaluation_scores(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取详细评分列表"""
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    scores = db.query(AIScore).filter(AIScore.project_id == project_id).all()
    
    return [
        AIScoreResponse(
            id=s.id,
            project_id=s.project_id,
            indicator_id=s.indicator_id,
            score=s.score,
            max_score=s.max_score,
            confidence=s.confidence,
            evidence_text=s.evidence_text,
            source_page=s.source_page,
            reasoning=s.reasoning,
            created_at=s.created_at
        )
        for s in scores
    ]


@router.post("/parse-document/{document_id}", response_model=MessageResponse)
async def parse_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """解析单个文档"""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    if not os.path.exists(document.file_url):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    result = await parser.parse(document.file_url, document.file_type or "pdf")
    
    if result.get("success"):
        document.parsed_text = result.get("text", "")
        document.extracted_data = result.get("extracted", {})
        document.ocr_quality = result.get("ocr_quality", 0)
        document.status = "parsed"
        db.commit()
        
        return MessageResponse(
            message="文档解析成功",
            data={
                "document_id": document_id,
                "text_length": len(document.parsed_text),
                "extracted_keys": list(document.extracted_data.keys()) if document.extracted_data else []
            }
        )
    else:
        document.status = "failed"
        db.commit()
        raise HTTPException(status_code=500, detail=f"文档解析失败: {result.get('error', '未知错误')}")


# ── 导出评审报告 ──────────────────────────────────────────

@router.get("/export/{project_id}/{format}")
async def export_evaluation_report(
    project_id: int,
    format: str,  # "word" or "pdf"
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    导出评审报告（Word/PDF）
    """
    from fastapi.responses import FileResponse
    import json
    
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    
    # 权限检查
    if current_user.role.value != "admin" and project.enterprise_id != current_user.id:
        raise HTTPException(status_code=403, detail="无权导出此项目报告")
    
    # 获取评审结果
    scores = db.query(AIScore).filter(AIScore.project_id == project_id).all()
    dimensions = _build_score_dimensions(scores)
    
    # 生成报告数据
    report_data = {
        "project_name": project.name,
        "domain": project.domain or "未指定",
        "total_score": project.total_score or 0,
        "grade": project.grade or "未评定",
        "dimensions": dimensions,
        "completed_at": project.completed_at.isoformat() if project.completed_at else None,
    }
    
    # 生成报告文件
    export_dir = "./data/exports"
    os.makedirs(export_dir, exist_ok=True)
    
    if format == "word":
        file_path = _generate_word_report(report_data, export_dir)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif format == "pdf":
        file_path = _generate_pdf_report(report_data, export_dir)
        media_type = "application/pdf"
    else:
        raise HTTPException(status_code=400, detail="不支持的格式，请使用 word 或 pdf")
    
    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=f"{project.name}_评审报告.{format}",
        headers={"Content-Disposition": f'attachment; filename="{project.name}_评审报告.{format}"'}
    )


def _generate_word_report(data: dict, export_dir: str) -> str:
    """生成 Word 报告"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # 如果没有 python-docx，生成简单的文本文件
        file_path = f"{export_dir}/{data['project_name']}_评审报告.docx"
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(f"绿色建筑先进适用技术评审报告\n\n")
            f.write(f"项目名称：{data['project_name']}\n")
            f.write(f"技术领域：{data['domain']}\n")
            f.write(f"总分：{data['total_score']}\n")
            f.write(f"等级：{data['grade']}\n\n")
            f.write("评审维度：\n")
            for dim in data['dimensions']:
                f.write(f"  - {dim['name']}: {dim['score']}/{dim['max_score']}\n")
        return file_path
    
    doc = Document()
    
    # 标题
    title = doc.add_heading('绿色建筑先进适用技术评审报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 项目信息
    doc.add_heading('一、项目基本信息', level=1)
    doc.add_paragraph(f"项目名称：{data['project_name']}")
    doc.add_paragraph(f"技术领域：{data['domain']}")
    doc.add_paragraph(f"评审日期：{data['completed_at'] or '未完成'}")
    
    # 评审结果
    doc.add_heading('二、评审结果', level=1)
    doc.add_paragraph(f"总分：{data['total_score']} 分")
    doc.add_paragraph(f"等级：{data['grade']}")
    
    # 各维度评分
    doc.add_heading('三、各维度评分详情', level=1)
    for dim in data['dimensions']:
        doc.add_paragraph(f"{dim['name']}: {dim['score']}/{dim['max_score']} 分")
        if dim.get('children'):
            for child in dim['children']:
                doc.add_paragraph(f"    - {child['name']}: {child['score']}/{child['max_score']} 分", style='List Bullet')
    
    file_path = f"{export_dir}/{data['project_name']}_评审报告.docx"
    doc.save(file_path)
    return file_path


def _generate_pdf_report(data: dict, export_dir: str) -> str:
    """生成 PDF 报告"""
    # 简单实现：生成文本文件作为 PDF 替代
    file_path = f"{export_dir}/{data['project_name']}_评审报告.pdf"
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write("=" * 50 + "\n")
        f.write("绿色建筑先进适用技术评审报告\n")
        f.write("=" * 50 + "\n\n")
        f.write(f"项目名称：{data['project_name']}\n")
        f.write(f"技术领域：{data['domain']}\n")
        f.write(f"总分：{data['total_score']}\n")
        f.write(f"等级：{data['grade']}\n\n")
        f.write("评审维度：\n")
        for dim in data['dimensions']:
            f.write(f"  - {dim['name']}: {dim['score']}/{dim['max_score']}\n")
    return file_path


# ── 后台任务 ──────────────────────────────────────────

async def _run_evaluation_task(project_id: int, user_id: int):
    """后台评审任务"""
    from database import SessionLocal
    
    db = SessionLocal()
    try:
        result = await evaluation_engine.run_full_evaluation(project_id, db, user_id)
        
        log = AuditLog(
            user_id=user_id,
            action="evaluation_completed" if result.get("success") else "evaluation_failed",
            module="evaluation",
            detail=f"项目 {project_id}: {result.get('message', result.get('grade', '失败'))}"
        )
        db.add(log)
        db.commit()
        
    except Exception as e:
        project = db.query(Project).filter(Project.id == project_id).first()
        if project:
            project.status = ProjectStatus.reviewing
        
        error_log = AuditLog(
            user_id=user_id,
            action="evaluation_failed",
            module="evaluation",
            detail=f"项目 {project_id} 评审失败: {str(e)}"
        )
        db.add(error_log)
        db.commit()
    finally:
        db.close()


# ── 辅助函数 ──────────────────────────────────────────

def _build_score_dimensions(scores: List[AIScore]) -> List[dict]:
    """构建评分维度树"""
    from models.indicators import get_all_indicators, get_l2_by_l1
    
    dimensions = []
    indicators = get_all_indicators()
    l1_indicators = [i for i in indicators if i.level.value == 1]
    
    for l1 in l1_indicators:
        l2_list = get_l2_by_l1(l1.id)
        l2_scores = []
        
        for l2 in l2_list:
            score_record = next((s for s in scores if s.indicator_id == l2.id), None)
            if score_record:
                l2_scores.append({
                    "indicator_id": l2.id,
                    "name": l2.name,
                    "score": score_record.score,
                    "max_score": l2.max_score,
                    "weight": l2.weight,
                    "confidence": score_record.confidence,
                    "evidence": score_record.evidence_text or "",
                    "reasoning": score_record.reasoning or ""
                })
        
        if l2_scores:
            l1_score = sum(s["score"] for s in l2_scores)
            l1_max = sum(s["max_score"] for s in l2_scores)
            
            dimensions.append({
                "indicator_id": l1.id,
                "name": l1.name,
                "score": l1_score,
                "max_score": l1_max,
                "weight": l1.weight,
                "confidence": sum(s["confidence"] for s in l2_scores) / len(l2_scores),
                "evidence": "",
                "reasoning": "",
                "children": l2_scores
            })
    
    return dimensions
