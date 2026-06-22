"""
文档解析服务 - PDF/Word/OCR
"""
import os
import tempfile
from typing import Dict, Any, Optional
from pathlib import Path


class DocumentParser:
    """文档解析引擎"""
    
    def __init__(self):
        self.supported_types = ["pdf", "docx", "doc", "txt", "md"]
    
    async def parse(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        解析文档，返回结构化数据
        """
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return await self._parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return await self._parse_docx(file_path)
        elif file_type in ["txt", "md"]:
            return await self._parse_text(file_path)
        else:
            return {"error": f"不支持的文件类型: {file_type}"}
    
    async def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析 PDF 文档"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(file_path)
            full_text = ""
            pages_text = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                pages_text.append({
                    "page": page_num + 1,
                    "text": text
                })
                full_text += text + "\n"
            
            doc.close()
            
            # 提取关键信息
            extracted = self._extract_key_info(full_text)
            
            return {
                "success": True,
                "text": full_text,
                "pages": pages_text,
                "page_count": len(pages_text),
                "extracted": extracted,
                "ocr_quality": 1.0  # PDF 直接提取，质量为 1
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析 Word 文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            full_text = ""
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
                    full_text += para.text + "\n"
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            extracted = self._extract_key_info(full_text)
            
            return {
                "success": True,
                "text": full_text,
                "paragraphs": paragraphs,
                "tables": tables,
                "extracted": extracted,
                "ocr_quality": 1.0
            }
            
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    async def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """解析纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read()
            
            extracted = self._extract_key_info(text)
            
            return {
                "success": True,
                "text": text,
                "extracted": extracted,
                "ocr_quality": 1.0
            }
        except Exception as e:
            return {"success": False, "error": str(e)}
    
    def _extract_key_info(self, text: str) -> Dict[str, Any]:
        """
        从文本中提取关键信息（供评审使用）
        """
        info = {
            "project_name": None,
            "company_name": None,
            "tech_description": None,
            "innovation_points": [],
            "energy_saving_rate": None,
            "carbon_reduction": None,
            "patents": [],
            "case_studies": [],
            "test_reports": []
        }
        
        # 简单的关键词提取（实际应该用 LLM 或 NLP）
        lines = text.split("\n")
        
        for i, line in enumerate(lines):
            line = line.strip()
            
            # 项目名称
            if "项目名称" in line or "申报项目" in line:
                if ":" in line or "：" in line:
                    info["project_name"] = line.split(":")[-1].split("：")[-1].strip()
            
            # 申报单位
            if "申报单位" in line or "单位名称" in line:
                if ":" in line or "：" in line:
                    info["company_name"] = line.split(":")[-1].split("：")[-1].strip()
            
            # 节能率
            if "节能率" in line and "%" in line:
                import re
                match = re.search(r"(\d+\.?\d*)%", line)
                if match:
                    info["energy_saving_rate"] = float(match.group(1))
            
            # 碳减排
            if "碳减排" in line and "%" in line:
                import re
                match = re.search(r"(\d+\.?\d*)%", line)
                if match:
                    info["carbon_reduction"] = float(match.group(1))
            
            # 专利
            if "专利" in line and ("发明" in line or "实用新型" in line):
                info["patents"].append(line)
            
            # 创新点
            if "创新点" in line or "技术创新" in line:
                # 收集后续几行
                for j in range(i+1, min(i+5, len(lines))):
                    if lines[j].strip() and not lines[j].strip().startswith("#"):
                        info["innovation_points"].append(lines[j].strip())
            
            # 案例
            if "案例" in line and ("项目" in line or "工程" in line):
                info["case_studies"].append(line)
            
            # 检测报告
            if "检测报告" in line or "检测单位" in line:
                info["test_reports"].append(line)
        
        return info


# 全局实例
parser = DocumentParser()
